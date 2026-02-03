from datetime import datetime, timedelta, timezone
import pandas as pd
import io
import openpyxl
from typing import List
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from Equipment.schemas import SEquipment, SEquipmentCreate, SEquipmentWithResponsible, SEquipmentExcelReport
from Equipment import crud
from User.depends import get_current_user
from User.models import User
from Equipment.reports import build_grouped_report, build_complex_report_all_categories

router = APIRouter(
    prefix="/equipment",
    tags=["Оборудование"]
)

@router.post("/to_excel_file")
async def generate_equipment_excel(payload: SEquipmentExcelReport, user: User = Depends(get_current_user)):
    if user.system_role_id < 2:
        raise HTTPException(status_code=403, detail="Forbidden")

    try:
        equipment_period = await crud.get_equipment_for_ready_report(payload.ids) \
            if hasattr(crud, "get_equipment_for_ready_report") else await crud.get_equipment_by_ids(payload.ids)

        #1: группировка
        if payload.report_type_id == 1:
            content = build_grouped_report(equipment_period)

        #2: сложносоставной
        elif payload.report_type_id == 2:
            categories = await crud.get_all_categories_with_types()

            categorized_type_ids: set[int] = set()
            for cat in categories:
                types = list(getattr(cat, "types", []) or [])
                for t in types:
                    categorized_type_ids.add(t.id)

            equipment_period = equipment_period
            period_type_counts: dict[int | None, int] = {}
            for e in equipment_period:
                tid = getattr(e, "type_id", None)
                period_type_counts[tid] = period_type_counts.get(tid, 0) + 1

            def _tname(t) -> str:
                return (getattr(t, "type_name", "") or getattr(t, "name", "") or "").strip()

            blocks = []

            for cat in categories:
                cat_name = getattr(cat, "category_name", "") or getattr(cat, "name", "") or ""
                types = list(getattr(cat, "types", []) or [])
                if not types:
                    continue

                type_ids = [t.id for t in types]
                totals_all = await crud.count_equipment_by_type_ids(type_ids)
                totals_period = await crud.count_equipment_by_type_ids_in_ids(type_ids, payload.ids)

                types_sorted = sorted(types, key=lambda t: _tname(t).lower())

                rows = []
                for t in types_sorted:
                    period_cnt = totals_period.get(t.id, 0)
                    if period_cnt == 0:
                        continue

                    rows.append((_tname(t), totals_all.get(t.id, 0), period_cnt))

                if rows:
                    blocks.append((cat_name, rows))

            other_type_ids: list[int] = []
            other_without_type = period_type_counts.get(None, 0)

            for tid, cnt in period_type_counts.items():
                if tid is None:
                    continue
                if tid not in categorized_type_ids:
                    other_type_ids.append(tid)

            if other_type_ids or other_without_type > 0:
                totals_all_other = await crud.count_equipment_by_type_ids(other_type_ids) if other_type_ids else {}
                totals_period_other = await crud.count_equipment_by_type_ids_in_ids(other_type_ids, payload.ids) if other_type_ids else {}

                type_id_to_name: dict[int, str] = {}
                for e in equipment_period:
                    tid = getattr(e, "type_id", None)
                    if tid is None or tid in type_id_to_name:
                        continue
                    t = getattr(e, "type", None)
                    if t is not None:
                        type_id_to_name[tid] = _tname(t)

                other_rows = []

                for tid in sorted(other_type_ids, key=lambda x: (type_id_to_name.get(x, ""), x)):
                    period_cnt = totals_period_other.get(tid, 0)
                    if period_cnt == 0:
                        continue
                    other_rows.append((
                        type_id_to_name.get(tid, f"Тип #{tid}"),
                        totals_all_other.get(tid, 0),
                        period_cnt
                    ))

                if other_without_type > 0:
                    other_rows.append(("Тип не указан", 0, other_without_type))

                if other_rows:
                    blocks.append(("Прочее", other_rows))

            content = build_complex_report_all_categories(blocks)
        else:
            raise HTTPException(status_code=400, detail="Unknown report_type_id")

        excel_file = io.BytesIO(content)
        excel_file.seek(0)

        file_name = f"report_{payload.report_type_id}_{datetime.now(tz=timezone(timedelta(hours=5))).strftime('%Y%m%d_%H%M%S')}.xlsx"

        return StreamingResponse(
            excel_file,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{file_name}"'}
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"Ошибка при генерации Excel-файла: {e}")
        raise HTTPException(status_code=500, detail="Ошибка при генерации Excel-файла")

@router.post("/create", response_model=SEquipment)
async def create_equipment(equipment: SEquipmentCreate):
    db_equipment = await crud.get_equipment_by_serial_number(serial_number=equipment.serial_number)
    if db_equipment:
        raise HTTPException(status_code=400, detail="Equipment with this serial number already exists")
    return await crud.create_equipment(equipment=equipment)

@router.get("/check_inventory/{inventory_number}")
async def check_inventory_number(inventory_number: str):
    equipment = await crud.get_equipment_by_inventory_number(inventory_number)
    if equipment:
        raise HTTPException(status_code=400, detail="Equipment with this inventory number already exists")

@router.get("/all")
async def get_all_equipment(user: User = Depends(get_current_user)) -> List[SEquipmentWithResponsible]:
    return await crud.get_all_equipment(user_role_id=user.system_role_id)

@router.get("/{equipment_id}", response_model=SEquipmentWithResponsible)
async def get_equipment(equipment_id: int):
    equipment = await crud.get_equipment(equipment_id)
    if not equipment:
        raise HTTPException(status_code=404, detail="Equipment not found")
    return equipment

@router.put("/{equipment_id}", response_model=SEquipment)
async def update_equipment(equipment_id: int, updated_equipment: SEquipmentCreate):
    existing_equipment = await crud.get_equipment(equipment_id)
    if not existing_equipment:
        raise HTTPException(status_code=404, detail="Equipment not found")
    
    return await crud.update_equipment(equipment_id=equipment_id, updated_equipment=updated_equipment)

@router.delete("/{equipment_id}", response_model=dict)
async def delete_equipment(equipment_id: int):
    return await crud.delete_equipment(equipment_id=equipment_id)

@router.get("/to_word/{equipment_id}")
async def generate_equipment_word(equipment_id: int, user: User = Depends(get_current_user)):
    equipment = await crud.get_equipment_for_word(equipment_id)
    if not equipment:
        raise HTTPException(status_code=404, detail="Equipment not found")

    doc = Document()
    
    def set_font(paragraph, size, bold=False, align=None):
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(size)
        font.bold = bold
        if align:
            paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(0)
    
    p1 = doc.add_paragraph('Национальный исследовательский университет "Высшая школа экономики"\nПермский филиал')
    set_font(p1, 12, bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    p2 = doc.add_paragraph("\nКАРТОЧКА\nУЧЕТА ОБОРУДОВАНИЯ\n")
    set_font(p2, 14, bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    fields = [
        ("Наименование:", equipment.type_name),
        ("Модель оборудования:", equipment.model),
        ("Серийный номер:", equipment.serial_number),
        ("Инвентарный номер:", equipment.inventory_number),
        ("Кому выдано:", equipment.responsible_user_full_name if equipment.responsible_user_full_name else "Не указано"),
        ("Подпись:", "_________________________/"),
        ("Подразделение:", equipment.responsible_user_office)
    ]
    
    for field_name, field_value in fields:
        p = doc.add_paragraph()
        run = p.add_run(field_name + " ")
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run2 = p.add_run(field_value)
        run2.font.name = "Times New Roman"
        run2.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
    
    now = datetime.now(tz=timezone(timedelta(hours=5))).strftime("%d/%m/%Y %I:%M:%S %p")
    p_last = doc.add_paragraph(now)
    set_font(p_last, 12, align=WD_PARAGRAPH_ALIGNMENT.RIGHT)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    file_name = f"{equipment.serial_number}.docx"
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={file_name}"}
    )